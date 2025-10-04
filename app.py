import streamlit as st
import time
import os
import shutil
import json
import glob
from io import BytesIO
from docx import Document
import re
from io import BytesIO
from datetime import datetime
import traceback


from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Preformatted
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor, grey
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib import colors
from main import (
    LocalSentenceTransformerEmbeddings,
    download_full_video_safe,
    extract_video_frames_optimized,
    get_transcript,
    translate_transcript_batch,
    _sanitize_metadata,
    cleanup_leftover_videos,
    cleanup_video_files_for_streamlit,
    cleanup_specific_file_types,
    detect_query_type,
    extract_video_id
)

from translation_and_chat import (
    process_video_and_qa, 
    answer_question,  
    load_existing_vectorstore,
    create_vector_store
)

from notes_generation import (
    process_video_for_notes,
    extract_topic_structure_with_flash,
    format_notes_for_display
)

st.set_page_config(page_title="U-Tube AI", layout="wide")

if 'video_ready' not in st.session_state:
    st.session_state.video_ready = False
if 'current_video_id' not in st.session_state:
    st.session_state.current_video_id = None
if 'notes_structure' not in st.session_state:
    st.session_state.notes_structure = None
if 'topic_outline' not in st.session_state:
    st.session_state.topic_outline = None
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []

def parse_markdown(md_text):
    blocks = []
    lines = md_text.strip().splitlines()
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code blocks
        if stripped.startswith("```"):
            code_block = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_block.append(lines[i])
                i += 1
            blocks.append(("code", "\n".join(code_block)))
            i += 1
            continue

        # Lists
        if stripped.startswith(('- ', '* ', '+ ')):
            list_items = []
            while i < len(lines) and lines[i].strip().startswith(('- ', '* ', '+ ')):
                list_items.append(lines[i].strip()[2:])
                i += 1
            blocks.append(("ul", list_items))
            continue
            
        if re.match(r'^\d+\.\s', stripped):
            list_items = []
            while i < len(lines) and re.match(r'^\d+\.\s', lines[i].strip()):
                list_items.append(re.sub(r'^\d+\.\s*', '', lines[i].strip()))
                i += 1
            blocks.append(("ol", list_items))
            continue
        
        # Tables
        if '|' in stripped:
            table = []
            header_check = [c.strip() for c in stripped.split('|') if c.strip()]
            if all(set(c) <= set('-: ') for c in header_check):
                i += 1
                continue
            
            while i < len(lines) and '|' in lines[i]:
                row = [c.strip() for c in lines[i].split('|')][1:-1]
                if all(set(c) <= set('-: ') for c in row):
                    i += 1
                    continue
                table.append(row)
                i += 1
            blocks.append(("table", table))
            continue

        # ✅ FIXED: Headings with safe parsing
        if stripped.startswith("#"):
            # Count heading level
            level = 0
            for char in stripped:
                if char == '#':
                    level += 1
                else:
                    break
            
            # Validate level (must be 1-6)
            if level > 0 and level <= 6:
                text = stripped[level:].strip()  # Remove # symbols
                blocks.append((f"h{level}", text))
            else:
                # Not a valid heading, treat as paragraph
                blocks.append(("p", stripped))
        
        elif stripped.startswith(">"):
            blocks.append(("quote", stripped[1:].strip()))
        elif stripped in ("---", "___", "***"):
            blocks.append(("hr", None))
        elif stripped:
            blocks.append(("p", stripped))
        
        i += 1

    return blocks


def _format_word_text(paragraph, text):
    pattern = r'(\[([^\]]+)\]\(([^\)]+)\))|(\*\*([^\*]+?)\*\*)|(__([^_]+?)__)|(\*([^\*]+?)\*)|(_([^_]+?)_)|(`([^`]+)`)'
    
    pos = 0
    for m in re.finditer(pattern, text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        
        if m.group(1):
            _add_word_hyperlink(paragraph, url=m.group(3), text=m.group(2))
        elif m.group(4):
            run = paragraph.add_run(m.group(5))
            run.bold = True
        elif m.group(6):
            run = paragraph.add_run(m.group(7))
            run.bold = True
        elif m.group(8):
            run = paragraph.add_run(m.group(9))
            run.italic = True
        elif m.group(10):
            run = paragraph.add_run(m.group(11))
            run.italic = True
        elif m.group(12):
            run = paragraph.add_run(m.group(13))
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        
        pos = m.end()

    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_word_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_paragraph_horizontal_rule(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)


def create_word_doc(markdown_text, title="Video Notes"):
    blocks = parse_markdown(markdown_text)
    
    doc = Document()
    doc.core_properties.title = title
    doc.core_properties.created = datetime.now()
    
    try:
        styles = doc.styles
        if 'Code Block' not in [s.name for s in styles]:
            code_style = styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
            code_style.font.name = 'Consolas'
            code_style.font.size = Pt(9)
    except Exception:
        pass

    for btype, content in blocks:
        try:  # ✅ Add error handling for each block
            if btype.startswith("h"):
                # ✅ Safe level extraction
                try:
                    level_str = btype[1:]  # Get everything after 'h'
                    if level_str.isdigit():
                        level = min(int(level_str), 4)
                    else:
                        print(f"Warning: Invalid heading level '{btype}', defaulting to 1")
                        level = 1
                except (ValueError, IndexError) as e:
                    print(f"Error parsing heading level from '{btype}': {e}")
                    level = 1
                
                doc.add_heading(content, level=level)
                
            elif btype == "p":
                para = doc.add_paragraph()
                _format_word_text(para, content)
            elif btype == "quote":
                para = doc.add_paragraph(style='Intense Quote')
                _format_word_text(para, content)
            elif btype == "ul":
                for item in content:
                    para = doc.add_paragraph(style="List Bullet")
                    _format_word_text(para, item)
            elif btype == "ol":
                for item in content:
                    para = doc.add_paragraph(style="List Number")
                    _format_word_text(para, item)
            elif btype == "code":
                try:
                    doc.add_paragraph(content, style='Code Block')
                except:
                    para = doc.add_paragraph(content)
                    for run in para.runs:
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
            elif btype == "hr":
                para = doc.add_paragraph()
                _add_paragraph_horizontal_rule(para)
            elif btype == "table" and content:
                num_rows = len(content)
                num_cols = max(len(r) for r in content) if content else 0
                if num_cols == 0:
                    continue
                
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Light Grid Accent 1'
                for i, row_data in enumerate(content):
                    for j in range(num_cols):
                        cell_text = row_data[j] if j < len(row_data) else ""
                        cell = table.cell(i, j)
                        cell.text = cell_text
                        if i == 0:
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    run.bold = True
        except Exception as e:
            print(f"Error processing block type '{btype}': {e}")
            continue  # Skip problematic blocks
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _format_pdf_text(text):
    text = re.sub(r'`([^`]+)`', r'<font name="Courier">\1</font>', text)
    text = re.sub(r'\[([^\]]+)\]\(([^\)]+)\)', r'<a href="\2" color="blue">\1</a>', text)
    text = re.sub(r'\*\*([^\*]+?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'__([^_]+?)__', r'<b>\1</b>', text)
    text = re.sub(r'(?<!\*)\*([^\*]+?)\*(?!\*)', r'<i>\1</i>', text)
    text = re.sub(r'(?<!_)_([^_]+?)_(?!_)', r'<i>\1</i>', text)
    return text


def create_pdf_doc(markdown_text, title="Video Notes"):
    blocks = parse_markdown(markdown_text)
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=inch, rightMargin=inch, topMargin=inch, bottomMargin=inch, title=title)
    styles = getSampleStyleSheet()
    
    quote_style = ParagraphStyle(name='CustomQuote', parent=styles['Normal'], fontName='Helvetica-Oblique', textColor=grey, leftIndent=20)
    code_style = ParagraphStyle(name='CustomCode', parent=styles['Normal'], fontName='Courier', fontSize=9, backColor=HexColor('#F0F0F0'), leftIndent=20, spaceBefore=6, spaceAfter=6)
    bullet_style = ParagraphStyle(name='CustomBullet', parent=styles['BodyText'], leftIndent=20, bulletIndent=10)
    
    flowables = []
    page_width = doc.width
    
    for btype, content in blocks:
        try:  # ✅ Add error handling for each block
            if btype.startswith("h"):
                # ✅ Safe level extraction
                try:
                    level_str = btype[1:]
                    if level_str.isdigit():
                        level = min(int(level_str), 3)
                    else:
                        print(f"Warning: Invalid heading level '{btype}', defaulting to 1")
                        level = 1
                except (ValueError, IndexError) as e:
                    print(f"Error parsing heading level from '{btype}': {e}")
                    level = 1
                
                style_name = f"Heading{level}"
                flowables.append(Paragraph(_format_pdf_text(content), styles[style_name]))
                flowables.append(Spacer(1, 12))
                
            elif btype == "p":
                flowables.append(Paragraph(_format_pdf_text(content), styles['BodyText']))
                flowables.append(Spacer(1, 8))
            elif btype == "quote":
                flowables.append(Paragraph(_format_pdf_text(content), quote_style))
                flowables.append(Spacer(1, 8))
            elif btype == "ul":
                for item in content:
                    flowables.append(Paragraph(f"• {_format_pdf_text(item)}", bullet_style))
                flowables.append(Spacer(1, 8))
            elif btype == "ol":
                for i, item in enumerate(content, 1):
                    flowables.append(Paragraph(f"{i}. {_format_pdf_text(item)}", bullet_style))
                flowables.append(Spacer(1, 8))
            elif btype == "code":
                flowables.append(Preformatted(content, code_style))
                flowables.append(Spacer(1, 12))
            elif btype == "hr":
                hr_table = Table([['']], colWidths=[page_width])
                hr_table.setStyle(TableStyle([('LINEBELOW', (0, 0), (-1, -1), 1, colors.grey)]))
                flowables.append(hr_table)
                flowables.append(Spacer(1, 12))
            elif btype == "table" and content:
                table = Table(content)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), HexColor('#4F81BD')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('BACKGROUND', (0, 1), (-1, -1), HexColor('#D3DFEE')),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('TOPPADDING', (0, 0), (-1, -1), 6),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ]))
                flowables.append(table)
                flowables.append(Spacer(1, 12))
        except Exception as e:
            print(f"Error processing block type '{btype}': {e}")
            continue  # Skip problematic blocks

    doc.build(flowables)
    buffer.seek(0)
    return buffer

def process_video_no_translation(youtube_url):
    """Process video without translation for English content"""
    try:
        print("Processing English video without translation...")
        
        video_id = extract_video_id(youtube_url)
        if not video_id:
            return None, "Invalid YouTube URL"
        
        # Check existing database
        db_path = f"./chroma_db_{video_id}"
        if os.path.exists(db_path):
            vectorstore = load_existing_vectorstore(video_id)
            if vectorstore:
                return vectorstore, "Database loaded successfully"
            else:
                shutil.rmtree(db_path, ignore_errors=True)
        
        # Extract transcript
        transcript_chunks = get_transcript(video_id)
        if not transcript_chunks:
            return None, "Failed to extract transcript"
        
        # Skip translation - use transcript directly
        print(f"Using {len(transcript_chunks)} transcript chunks directly (no translation)")
        
        # Create vector store directly
        vectorstore = create_vector_store(transcript_chunks, video_id)
        if not vectorstore:
            return None, "Failed to create vector store"
        
        return vectorstore, "Video processed successfully (no translation)"
        
    except Exception as e:
        return None, f"Error processing video: {str(e)}"

with st.sidebar:
    st.title("U-Tube AI App")
    st.markdown("## About")
    st.markdown("""
    U-Tube AI App allows users to interact with video content using AI.
    Users can upload videos, ask questions about the content, and receive answers and Generate Notes powered by advanced AI models.
    """)
    
    st.markdown("---")

    st.markdown("## How to Use")
    st.markdown("""
    1. Upload a YouTube video by entering its URL.
    2. Select the desired language for interaction.
    3. Choose a task: Chat with Video, Important Points, Make Notes, or Translate Video.
    4. Click Submit to process the video.
    5. Press "Clear and Resrt" to start over with a new video or to switch to other tasks of the same video.
    6. Use the "Force Rebuild Database" button to delete all existing databases and start fresh if you came up with an error or needed.
    7. If the error persists, please click the "Clear and Restart" button to reset the app state.
    8. If error still persists, please check the console logs for more details.
    9. Thank you for using U-Tube AI App! Please star the GitHub repo if you find it useful and share with others.
    """)
    
    st.markdown("---")
    
    youtube_url = st.text_input("YouTube URL", "")
    language = st.selectbox("Select Language", ["English", "Hindi", "Spanish", "French", "German"])
    task_option = st.radio(
        "What do you want to do?",
        ("Chat with Video", "Important Points", "Make Notes", "Translate Video")
    )
    st.markdown("---")
    st.markdown("### 📸 Notes Options")
    include_visuals = st.radio(
        "Add visual information to notes?",
        ("Yes - Include video screenshots & OCR (It will take longer time, so press Yes only if necessary Visual content present in the video.)", "No - Text only"),
        help="Visual info includes screenshots, diagrams, and text from images. Select decisively as it affects processing time and cost."
    )
    
    # Store in session state
    st.session_state.include_visuals = (include_visuals == "Yes - Include screenshots & OCR, it will take longer than usual")
    submit_button = st.button("Submit")
    
    if st.button("🔄 Force Rebuild Database"):
        if youtube_url.strip():
            video_id = extract_video_id(youtube_url)
            if video_id:
                cleanup_video_files_for_streamlit(video_id)
                
                import glob
                for db_folder in glob.glob("./chroma_db_*"):
                    if os.path.exists(db_folder):
                        shutil.rmtree(db_folder, ignore_errors=True)
                
                st.success("🗑️ All databases deleted! Click Submit to rebuild.")
    
    st.markdown("---")
    
    if st.session_state.video_ready:
        st.success(f"✅ Video Ready: {st.session_state.current_video_id}")
        st.info(f"📋 Selected Language: {language}")
    else:
        st.warning("⏳ No video processed")

st.title("YouTube Content LM Synthesiser")
st.markdown("### Upload a YouTube video, generate notes and interact with its content using Gemini.")

if submit_button:
    if youtube_url.strip():
        
        video_id = extract_video_id(youtube_url)
        if not video_id:
            st.error("❌ Invalid YouTube URL")
            st.stop()
        
        # Clear chat history for new video
        st.session_state.chat_history = []
        
        # Delete the specific video database
        db_path = f"./chroma_db_{video_id}"
        if os.path.exists(db_path):
            shutil.rmtree(db_path, ignore_errors=True)

        # Also delete any ChromaDB cache/metadata files
        for db_folder in glob.glob("./chroma_db_*"):
            if os.path.exists(db_folder):
                shutil.rmtree(db_folder, ignore_errors=True)

        # Delete any .chroma files or directories in current folder
        if os.path.exists("./.chroma"):
            shutil.rmtree("./.chroma", ignore_errors=True)
            time.sleep(1)
        
        with st.spinner("Processing video..."):
            try:
                if language == "English":
                    st.info("📝 Processing English video (skipping translation)")
                    vectorstore, result = process_video_no_translation(youtube_url)
                else:
                    st.info(f"🌍 Processing video with translation to {language}")
                    vectorstore, result = process_video_and_qa(youtube_url)
                
                if vectorstore and "successfully" in str(result):
                    test_search = vectorstore.similarity_search("video", k=1)
                    
                    if test_search:
                        st.session_state.current_video_id = video_id
                        st.session_state.video_ready = True
                        
                        st.success("✅ Video processed successfully!")
                        st.success(f"🔍 Database created with searchable content")
                        st.info(f"📊 Result: {result}")
                        st.rerun()
                    else:
                        st.error("❌ Database created but has no searchable content!")
                else:
                    st.error(f"❌ Processing failed: {result}")
                    
            except Exception as e:
                st.error(f"❌ Error: {str(e)}")
    else:
        st.warning("⚠️ Please enter a YouTube URL")

# Task handling
if st.session_state.video_ready:
    st.markdown("---")
    
    if task_option == "Chat with Video":
        st.markdown("### 💬 Chat with Video")
        
        # Add clear chat button
        col1, col2 = st.columns([6, 1])
        with col2:
            if st.button("🗑️ Clear Chat"):
                st.session_state.chat_history = []
                st.rerun()
        
        # Display chat history
        for message in st.session_state.chat_history:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
                
                # Display additional metadata for assistant messages
                if message["role"] == "assistant" and "metadata" in message:
                    metadata = message["metadata"]
                    if metadata.get('confidence'):
                        st.caption(f"🎯 Confidence: {metadata['confidence']:.2f}")
                    if metadata.get('timestamps'):
                        st.caption(f"⏱️ Timestamps: {metadata['timestamps']}")
        
        # Chat input
        user_input = st.chat_input("Ask me anything about the video...")
        
        if user_input:
            # Add user message to chat history
            st.session_state.chat_history.append({
                "role": "user",
                "content": user_input
            })
            
            # Display user message
            with st.chat_message("user"):
                st.markdown(user_input)
            
            # Generate assistant response
            with st.chat_message("assistant"):
                with st.spinner("Thinking..."):
                    try:
                        vectorstore = load_existing_vectorstore(st.session_state.current_video_id)
                        
                        if vectorstore:
                            result = answer_question(
                                vectorstore, 
                                st.session_state.current_video_id, 
                                user_input
                            )
                            
                            # Display answer
                            st.markdown(result['answer'])
                            
                            # Display metadata
                            if result.get('confidence'):
                                st.caption(f"🎯 Confidence: {result['confidence']:.2f}")
                            if result.get('timestamps'):
                                st.caption(f"⏱️ Timestamps: {result['timestamps']}")
                            
                            # Add assistant message to chat history
                            st.session_state.chat_history.append({
                                "role": "assistant",
                                "content": result['answer'],
                                "metadata": {
                                    "confidence": result.get('confidence'),
                                    "timestamps": result.get('timestamps'),
                                    "visual_elements": result.get('visual_elements'),
                                    "ocr_text": result.get('ocr_text')
                                }
                            })
                        else:
                            error_msg = "❌ Could not load video data. Please try reprocessing the video."
                            st.error(error_msg)
                            st.session_state.chat_history.append({
                                "role": "assistant",
                                "content": error_msg
                            })
                            
                    except Exception as e:
                        error_msg = f"❌ Error: {str(e)}"
                        st.error(error_msg)
                        st.session_state.chat_history.append({
                            "role": "assistant",
                            "content": error_msg
                        })
    
    elif task_option == "Important Points":
        st.markdown("### ⭐ Important Points & Topics")
        
        if st.button("🔍 Extract Important Points"):
            with st.spinner("Analyzing video structure... This may take 10-30 seconds"):
                try:
                    # Get transcript
                    transcript_chunks = get_transcript(st.session_state.current_video_id)
                    
                    if transcript_chunks:
                        # Extract topic structure using Flash
                        topic_structure = extract_topic_structure_with_flash(transcript_chunks)
                        st.session_state.topic_outline = topic_structure
                        
                        # Display results
                        st.success("✅ Important points extracted!")
                        
                        st.markdown(f"## {topic_structure.get('video_title', 'Video Summary')}")
                        st.info(f"**Duration:** {topic_structure.get('total_duration', 'Unknown')}")
                        
                        # Display topics
                        for topic in topic_structure.get('topics', []):
                            with st.expander(f"📌 Topic {topic.get('topic_id')}: {topic.get('title')}", expanded=True):
                                st.markdown(f"**Time Range:** {topic.get('start_time')} - {topic.get('end_time')}")
                                
                                st.markdown("**Subtopics:**")
                                for subtopic in topic.get('subtopics', []):
                                    st.markdown(f"- **{subtopic.get('subtopic_id')}. {subtopic.get('title')}**")
                                    st.markdown(f"  - *Time:* {subtopic.get('start_time')} - {subtopic.get('end_time')}")
                                    st.markdown(f"  - *Type:* {subtopic.get('content_type', 'N/A')}")
                                    st.markdown(f"  - *Complexity:* {subtopic.get('complexity', 'N/A')}")
                                    st.markdown(f"  - *Description:* {subtopic.get('description', 'N/A')}")
                                    st.markdown("")
                        
                        # Download options for topic outline
                        col1, col2 = st.columns(2)
                        with col1:
                            json_str = json.dumps(topic_structure, indent=2, ensure_ascii=False)
                            st.download_button(
                                label="📥 Download as JSON",
                                data=json_str,
                                file_name=f"topics_{st.session_state.current_video_id}.json",
                                mime="application/json"
                            )
                        
                    else:
                        st.error("Failed to extract transcript")
                        
                except Exception as e:
                    st.error(f"Error: {str(e)}")
    
    elif task_option == "Make Notes":
        st.markdown("### 📝 Comprehensive Notes Generation")
        
        # Visual info toggle
        col1, col2 = st.columns([3, 1])
        with col1:
            include_visuals = st.toggle(
                "📸 Include visual information (screenshots & OCR)",
                value=True,
                help="Adds screenshots, diagrams, and text from images to notes. Takes longer but provides richer content."
            )
        with col2:
            if include_visuals:
                st.caption("🎨 Enhanced Mode")
            else:
                st.caption("⚡ Fast Mode")
        
        # Store in session state
        st.session_state.include_visuals = include_visuals
        
        # Show processing info
        if include_visuals:
            st.info("⚠️ Enhanced mode: This takes a few minutes as it processes visual elements with advanced AI")
        else:
            st.info("⚡ Fast mode: Text-only processing - quicker but without visual context")
        
        if st.button("📝 Generate Comprehensive Notes"):
            with st.spinner("Generating comprehensive notes... Please wait."):
                try:
                    # Pass include_visuals parameter
                    success, formatted_notes, notes_structure = process_video_for_notes(
                        f"https://www.youtube.com/watch?v={st.session_state.current_video_id}",
                        include_visuals=st.session_state.include_visuals
                    )
                    
                    if success:
                        st.session_state.notes_structure = notes_structure
                        
                        st.success("✅ Notes generated successfully!")
                        
                        # Display statistics
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            st.metric("Topics", len(notes_structure.get('topics', [])))
                        with col2:
                            total_subtopics = sum(len(t.get('subtopics', [])) for t in notes_structure.get('topics', []))
                            st.metric("Subtopics", total_subtopics)
                        with col3:
                            if st.session_state.include_visuals:
                                visual_count = notes_structure.get('visual_elements_count', 0)
                                st.metric("Visual Elements", visual_count)
                            else:
                                st.metric("Mode", "Text Only")
                        
                        # Show processing mode badge
                        if st.session_state.include_visuals:
                            st.success("📸 Visual information included in notes")
                        else:
                            st.info("📝 Text-only notes generated")
                        
                        # Display notes
                        st.markdown("---")
                        st.markdown(formatted_notes)
                        
                        # Download options
                        st.markdown("---")
                        st.markdown("### 📥 Download Notes")
                        
                        col1, col2, col3 = st.columns(3)
                        
                        with col1:
                            # Markdown download
                            st.download_button(
                                label="📄 Download Markdown",
                                data=formatted_notes,
                                file_name=f"notes_{st.session_state.current_video_id}.md",
                                mime="text/markdown"
                            )
                        
                        with col2:
                            # Word download
                            try:
                                word_buffer = create_word_doc(formatted_notes, notes_structure.get('video_title', 'Video Notes'))
                                st.download_button(
                                    label="📘 Download Word",
                                    data=word_buffer,
                                    file_name=f"notes_{st.session_state.current_video_id}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                            except Exception as e:
                                st.error(f"Word generation error: {e}")
                        
                        with col3:
                            # PDF download
                            try:
                                pdf_buffer = create_pdf_doc(formatted_notes, notes_structure.get('video_title', 'Video Notes'))
                                st.download_button(
                                    label="📕 Download PDF",
                                    data=pdf_buffer,
                                    file_name=f"notes_{st.session_state.current_video_id}.pdf",
                                    mime="application/pdf"
                                )
                            except Exception as e:
                                st.error(f"PDF generation error: {e}")
                        
                        # Show processing summary
                        st.markdown("---")
                        st.markdown("#### 📊 Processing Summary")
                        summary_col1, summary_col2 = st.columns(2)
                        with summary_col1:
                            st.write(f"**Processing Mode:** {'Enhanced (with visuals)' if st.session_state.include_visuals else 'Fast (text-only)'}")
                            st.write(f"**Video Duration:** {notes_structure.get('total_duration', 'Unknown')}")
                        with summary_col2:
                            st.write(f"**Total Topics:** {len(notes_structure.get('topics', []))}")
                            st.write(f"**Total Subtopics:** {sum(len(t.get('subtopics', [])) for t in notes_structure.get('topics', []))}")
                            if st.session_state.include_visuals:
                                st.write(f"**Visual Elements:** {notes_structure.get('visual_elements_count', 0)}")
                        
                    else:
                        st.error(f"Failed to generate notes: {formatted_notes}")
                        
                except Exception as e:
                    st.error(f"Error: {str(e)}")
                    with st.expander("🔍 View Error Details"):
                        st.code(traceback.format_exc())
    
    elif task_option == "Translate Video":
        st.markdown("### 🌍 Translate Video")
        
        if st.session_state.video_ready:
            st.success("✅ Video transcript processed and ready!")
            st.info(f"**Target Language:** {language}")
            
            if st.button(f"Translate to {language}"):
                if language == "English":
                    st.success("✅ Content is already processed in English!")
                else:
                    st.info(f"🚧 Translation to {language} coming soon!")
        else:
            st.warning("⚠️ Please process a video first.")

# Cleanup
if st.session_state.video_ready:
    st.markdown("---")
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("🧹 Clear & Restart"):
            try:
                cleanup_video_files_for_streamlit(st.session_state.current_video_id)
                
                db_path = f"./chroma_db_{st.session_state.current_video_id}"
                if os.path.exists(db_path):
                    shutil.rmtree(db_path, ignore_errors=True)
            except:
                pass
            
            st.session_state.video_ready = False
            st.session_state.current_video_id = None
            st.session_state.notes_structure = None
            st.session_state.topic_outline = None
            st.session_state.chat_history = []
            st.rerun()
    
    with col2:
        if st.button("🚪 Exit"):
            st.success("👋 Thanks for using VidChat AI!")
            st.stop()

st.markdown("---")
st.markdown("*Built by Rishab © 2025. All rights reserved. Contact me if you found any issues or need help.*")